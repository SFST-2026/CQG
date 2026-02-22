#!/usr/bin/env python3
"""Update provenance metadata across the repo.

Single command to populate commit/digest/seed consistently in:
  - plot_metadata.json
  - outputs/provenance.{json,txt}
  - manuscript.docx placeholders (c18878630415a7e3598158d560cf7e615dba8967bdc8283472a610a7e362bedc, c04378d8b099, 123456, ...)

Typical use:
  make provenance DOCKER_IMAGE=sfst-qfis:local
  # or without docker:
  make provenance DOCKER_DIGEST_SHA256=sha256:... 
"""

from __future__ import annotations

import argparse
import json
import os
import re
import subprocess
from datetime import datetime, timezone
from pathlib import Path

import yaml


REPO_ROOT = Path(__file__).resolve().parents[1]


def _run(cmd: list[str]) -> str:
    return subprocess.check_output(cmd, cwd=REPO_ROOT, stderr=subprocess.STDOUT, text=True).strip()


def get_git_commit() -> str:
    env_commit = os.getenv("GIT_COMMIT")
    if env_commit:
        return env_commit.strip()
    try:
        return _run(["git", "rev-parse", "HEAD"])
    except Exception:
        return "UNKNOWN"


def get_seed(seeds_path: Path) -> str:
    if not seeds_path.exists():
        return os.getenv("SEED", "UNKNOWN")
    data = yaml.safe_load(seeds_path.read_text(encoding="utf-8")) or {}
    for key in ("global_seed", "numpy_seed", "python_random_seed"):
        if key in data and data[key] is not None:
            return str(data[key])
    return os.getenv("SEED", "UNKNOWN")


def get_docker_digest(docker_image: str | None) -> str:
    env_digest = os.getenv("DOCKER_DIGEST_SHA256")
    if env_digest:
        return env_digest.strip()

    docker_image = docker_image or os.getenv("DOCKER_IMAGE")
    if not docker_image:
        return "UNKNOWN"

    try:
        out = _run(["docker", "image", "inspect", docker_image, "--format", "{{json .RepoDigests}}"])
        digests = json.loads(out)
        if isinstance(digests, list):
            for d in digests:
                if "sha256:" in d:
                    return d.split("@")[-1]
    except Exception:
        pass
    return "UNKNOWN"


def update_plot_metadata(plot_metadata_path: Path, template_path: Path, payload: dict) -> None:
    base: dict = {}
    if template_path.exists():
        try:
            base = json.loads(template_path.read_text(encoding="utf-8"))
        except Exception:
            base = {}
    if plot_metadata_path.exists():
        try:
            base.update(json.loads(plot_metadata_path.read_text(encoding="utf-8")))
        except Exception:
            pass
    base.update(payload)
    plot_metadata_path.write_text(json.dumps(base, indent=2, sort_keys=True) + "\n", encoding="utf-8")


def write_outputs(outputs_dir: Path, payload: dict) -> None:
    outputs_dir.mkdir(parents=True, exist_ok=True)
    (outputs_dir / "provenance.json").write_text(json.dumps(payload, indent=2, sort_keys=True) + "\n", encoding="utf-8")
    lines = [
        f"timestamp_utc: {payload.get('timestamp_utc','')}",
        f"git_commit: {payload.get('git_commit','')}",
        f"docker_digest_sha256: {payload.get('docker_digest_sha256','')}",
        f"seed: {payload.get('seed','')}",
    ]
    (outputs_dir / "provenance.txt").write_text("\n".join(lines) + "\n", encoding="utf-8")


def replace_placeholders_in_docx(docx_path: Path, replacements: dict[str, str]) -> bool:
    try:
        from docx import Document
    except Exception as e:
        raise RuntimeError("python-docx is required") from e

    if not docx_path.exists():
        return False
    doc = Document(str(docx_path))
    modified = False

    def _process_runs(runs):
        nonlocal modified
        for r in runs:
            t = r.text
            new_t = t
            for k, v in replacements.items():
                if k in new_t:
                    new_t = new_t.replace(k, v)
            if new_t != t:
                r.text = new_t
                modified = True

    for p in doc.paragraphs:
        _process_runs(p.runs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _process_runs(p.runs)

    if modified:
        doc.save(str(docx_path))
    return modified




AUTO_BLOCK_START = "<!-- AUTO-PROVENANCE-START -->"
AUTO_BLOCK_END = "<!-- AUTO-PROVENANCE-END -->"


def _stamp_markdown_file(path: Path, payload: dict) -> bool:
    """Insert or replace an auto-generated provenance block in a markdown file."""
    if not path.exists():
        return False
    txt = path.read_text(encoding="utf-8", errors="ignore")
    block = "\n".join([
        AUTO_BLOCK_START,
        f"- git commit: `{payload.get('git_commit','')}`",
        f"- docker digest: `{payload.get('docker_digest_sha256','')}`",
        f"- seed: `{payload.get('seed','')}`",
        f"- timestamp (UTC): `{payload.get('timestamp_utc','')}`",
        AUTO_BLOCK_END,
        ""
    ])
    if AUTO_BLOCK_START in txt and AUTO_BLOCK_END in txt:
        new_txt = re.sub(re.escape(AUTO_BLOCK_START) + r".*?" + re.escape(AUTO_BLOCK_END),
                         block.strip(), txt, flags=re.DOTALL)
    else:
        # insert after first heading if present, otherwise prepend
        m = re.search(r"^# .*$", txt, flags=re.MULTILINE)
        if m:
            pos = m.end()
            new_txt = txt[:pos] + "\n\n" + block + txt[pos:]
        else:
            new_txt = block + txt
        # Replace placeholder tokens (<<...>>) in the markdown template if present
    for k,v in {
        'c18878630415a7e3598158d560cf7e615dba8967bdc8283472a610a7e362bedc': payload.get('docker_digest_sha256','').replace('sha256:',''),
        'c04378d8b099': payload.get('git_commit',''),
        '123456': str(payload.get('seed','')),
    }.items():
        new_txt = new_txt.replace(k, v)

if new_txt != txt:
        path.write_text(new_txt, encoding="utf-8")
        return True
    return False
def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--docker-image", default=None)
    ap.add_argument("--manuscript", default="manuscript.docx")
    args = ap.parse_args()

    payload = {
        "timestamp_utc": datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"),
        "git_commit": get_git_commit(),
        "docker_digest_sha256": get_docker_digest(args.docker_image),
        "seed": get_seed(REPO_ROOT / "config" / "seeds.yaml"),
    }

    update_plot_metadata(REPO_ROOT / "plot_metadata.json", REPO_ROOT / "plot_metadata.json.template", payload)
    write_outputs(REPO_ROOT / "outputs", payload)

    replacements = {
        "c18878630415a7e3598158d560cf7e615dba8967bdc8283472a610a7e362bedc": payload["docker_digest_sha256"],
        "<docker_digest>": payload["docker_digest_sha256"],
        "<docker_digest_sha256>": payload["docker_digest_sha256"],
        "c18878630415a7e3598158d560cf7e615dba8967bdc8283472a610a7e362bedc": payload["docker_digest_sha256"].replace("sha256:",""),
        "c04378d8b099": payload["git_commit"],
        "123456": payload["seed"],
        "c04378d8b099": payload["git_commit"],
        "<git_commit>": payload["git_commit"],
        "123456": payload["seed"],
    }

replace_placeholders_in_docx(REPO_ROOT / args.manuscript, replacements)

    # Stamp provenance into human-readable markdown docs (optional but reviewer-friendly)
    _stamp_markdown_file(REPO_ROOT / "supplement" / "SUPPLEMENT_FIGX_PROVENANCE.md", payload)
    _stamp_markdown_file(REPO_ROOT / "PROVENANCE.md", payload)

    print(json.dumps(payload, indent=2))


if __name__ == "__main__":
    main()
